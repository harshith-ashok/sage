"""Renders a small, deliberately limited Markdown subset (headings, **bold**,
*italic*, `code`, bullet/numbered lists, plain paragraphs, GFM pipe tables,
and LaTeX math) into real python-docx formatting — used by app.agent's
draft_docx_approval_note tool so the model can write Markdown (the same
thing it's told to write everywhere else) and get an actual formatted Word
document instead of literal ** characters, a table dumped as one run-on line
of pipe characters, or raw \\[...\\]/$$...$$ LaTeX source.

Tables become real Word tables (docx.table.Table), not text. Math is
rendered via docx_math.py (matplotlib's mathtext) and inserted as a small
inline/centered image — not a native, in-Word-editable equation object (that
would need converting LaTeX to OOXML's OMML format, which in turn needs
either a full LaTeX/MathML→OMML converter or Microsoft's own XSLT of
uncertain redistribution terms for an offline tool; rendering to an image
avoids that entirely and still reads as a properly typeset equation). A
single equation mathtext can't parse degrades to showing its raw LaTeX
source as text rather than aborting the whole document.
"""

import io
import re

from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Emu, Pt

from app.tasks.docx_math import LatexRenderError, render_latex_png

_HEADING_RE = re.compile(r"^(#{1,3})\s+(.*)")
_BULLET_RE = re.compile(r"^\s*[-*]\s+")
_NUMBERED_RE = re.compile(r"^\s*\d+\.\s+")
_TABLE_SEP_CELL_RE = re.compile(r"^:?-{1,}:?$")

# Bold/italic/code/math, matched in one left-to-right pass so they interleave
# correctly (e.g. "**bold with $x^2$ math**") instead of one pass per marker
# type clobbering the others' matches.
_INLINE_TOKEN_RE = re.compile(
    r"\*\*(?P<bold>.+?)\*\*"
    r"|\*(?P<italic>.+?)\*"
    r"|`(?P<code>.+?)`"
    r"|\$\$(?P<display_math>.+?)\$\$"
    r"|\$(?P<inline_math>[^\n$]+?)\$",
    re.DOTALL,
)

# The model reliably writes \[...\]/\(...\) (confirmed live — see
# frontend/src/lib/markdown.ts's normalizeLatexDelimiters, which has the same
# fix for the Console) rather than the $/$$ this module's tokenizer above
# looks for, so both are normalized to $ delimiters up front. Skips inline
# `code` spans so a literal "\(" inside one isn't touched.
_CODE_SPAN_RE = re.compile(r"`[^`\n]*`")


def _normalize_latex_delimiters(text: str) -> str:
    def convert(segment: str) -> str:
        segment = re.sub(r"\\\[([\s\S]+?)\\\]", lambda m: f"$${m.group(1)}$$", segment)
        segment = re.sub(r"\\\(([\s\S]+?)\\\)", lambda m: f"${m.group(1)}$", segment)
        return segment

    parts = _CODE_SPAN_RE.split(text)
    spans = _CODE_SPAN_RE.findall(text)
    out = []
    for i, part in enumerate(parts):
        out.append(convert(part))
        if i < len(spans):
            out.append(spans[i])
    return "".join(out)


def _add_equation_run(paragraph, latex: str, *, block: bool) -> None:
    """Inserts a rendered LaTeX image as a run in `paragraph`. Falls back to
    the raw "$latex$" text (still useful, just unrendered) if mathtext can't
    parse it, rather than dropping the content or crashing the document."""
    try:
        png_bytes, aspect = render_latex_png(latex, fontsize=15 if block else 12)
    except LatexRenderError:
        run = paragraph.add_run(f"${latex}$" if not block else f"$${latex}$$")
        run.font.name = "Courier New"
        return
    target_height = Pt(24) if block else Pt(12)
    target_width = Emu(int(int(target_height) * aspect))
    max_width = Emu(int(Pt(468)))  # ~6.5in, inside a standard letter-page margin
    if target_width > max_width:
        target_width = max_width
        target_height = Emu(int(int(target_width) / aspect))
    paragraph.add_run().add_picture(io.BytesIO(png_bytes), width=target_width, height=target_height)


def _add_inline_runs(paragraph, text: str) -> None:
    """Splits `text` on **bold**, *italic*, `code`, and $math$ markers,
    adding each as a separately-formatted run (or, for math, a rendered
    equation image) instead of one plain run with literal marker characters
    in it."""
    pos = 0
    for m in _INLINE_TOKEN_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        if m.group("bold") is not None:
            paragraph.add_run(m.group("bold")).bold = True
        elif m.group("italic") is not None:
            paragraph.add_run(m.group("italic")).italic = True
        elif m.group("code") is not None:
            run = paragraph.add_run(m.group("code"))
            run.font.name = "Courier New"
        elif m.group("display_math") is not None:
            _add_equation_run(paragraph, m.group("display_math"), block=False)
        else:
            _add_equation_run(paragraph, m.group("inline_math"), block=False)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _split_table_row(line: str) -> list[str]:
    """Splits a pipe-delimited row on "|", except a "|" inside a $math$/
    $$math$$/`code` span, which is content (e.g. \\left| x \\right| for
    absolute value), not a column separator. Caught live: a naive
    line.split("|") corrupted both the table structure and the LaTeX itself
    whenever an equation used a literal pipe, leaving fragments like
    "$\\displaystyle" as an orphaned "cell" with the rest of the expression
    sheared off into neighboring cells."""
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    cells: list[str] = []
    current: list[str] = []
    in_math = False
    in_code = False
    i, n = 0, len(stripped)
    while i < n:
        if not in_code and stripped[i : i + 2] == "$$":
            in_math = not in_math
            current.append("$$")
            i += 2
            continue
        ch = stripped[i]
        if ch == "`" and not in_math:
            in_code = not in_code
            current.append(ch)
        elif ch == "$" and not in_code:
            in_math = not in_math
            current.append(ch)
        elif ch == "|" and not in_math and not in_code:
            cells.append("".join(current).strip())
            current = []
        else:
            current.append(ch)
        i += 1
    cells.append("".join(current).strip())
    return cells


def _looks_like_table(lines: list[str]) -> bool:
    if len(lines) < 2 or "|" not in lines[0]:
        return False
    sep_cells = _split_table_row(lines[1])
    return bool(sep_cells) and all(_TABLE_SEP_CELL_RE.match(c) for c in sep_cells)


def _render_table(doc: DocumentType, lines: list[str]) -> None:
    header = _split_table_row(lines[0])
    data_rows = [_split_table_row(line) for line in lines[2:] if line.strip()]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    for cell, text in zip(table.rows[0].cells, header):
        p = cell.paragraphs[0]
        _add_inline_runs(p, text)
        for run in p.runs:
            run.bold = True
    for row_values in data_rows:
        cells = table.add_row().cells
        for i, cell in enumerate(cells):
            _add_inline_runs(cell.paragraphs[0], row_values[i] if i < len(row_values) else "")


def write_markdown(doc: DocumentType, markdown_text: str) -> None:
    """Appends `markdown_text` to `doc` as real Word formatting."""
    markdown_text = _normalize_latex_delimiters(markdown_text)
    for block in markdown_text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        lines = block.split("\n")

        heading_match = _HEADING_RE.match(lines[0])
        if heading_match and len(lines) == 1:
            doc.add_heading(heading_match.group(2), level=len(heading_match.group(1)))
            continue

        if _looks_like_table(lines):
            _render_table(doc, lines)
            continue

        if all(_BULLET_RE.match(line) or _NUMBERED_RE.match(line) for line in lines):
            for line in lines:
                is_bullet = _BULLET_RE.match(line) is not None
                content = _BULLET_RE.sub("", line) if is_bullet else _NUMBERED_RE.sub("", line)
                p = doc.add_paragraph(style="List Bullet" if is_bullet else "List Number")
                _add_inline_runs(p, content)
            continue

        # A block that's nothing but one display-math span ("$$...$$" alone,
        # e.g. a fitted-equation callout) gets its own centered, larger
        # rendering instead of being squeezed into inline size like math
        # embedded mid-sentence.
        stripped = block.strip()
        display_match = re.fullmatch(r"\$\$([\s\S]+?)\$\$", stripped)
        if display_match:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_equation_run(p, display_match.group(1), block=True)
            continue

        p = doc.add_paragraph()
        _add_inline_runs(p, " ".join(lines))
