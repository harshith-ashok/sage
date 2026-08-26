"""Phase 4: agentic document task. A scanned/photographed inspection report
(image) goes in; a grounded, SOP-checked .docx approval note comes out.

Pipeline (LangGraph StateGraph, same plan -> act -> verify shape as Phase
1's app/orchestrator.py, specialized here into concrete nodes):

    ocr -> draft -> review -(rejected, attempts left)-> draft
                       \\-(approved, or attempts exhausted)-> finalize

- ocr: vision model (app.router "vision" task type) reads the image and
  transcribes its findings as plain text.
- draft: reasoning model drafts an approval note, grounded in SOP context
  retrieved from Phase 3's knowledge base (queried using the OCR'd
  findings) — every requirement or decision must cite (Page X, Section Y).
- review: an independent reasoning-model call acts as compliance
  reviewer/critic, checking the draft against the same SOP context —
  catches a draft that just parrots a technician's recommendation instead
  of applying the SOP's actual escalation rule. Rejection feeds back into
  another draft attempt (bounded by MAX_ATTEMPTS).
- finalize: writes the accepted draft to a real .docx file via python-docx.
"""

import os
import uuid
from typing import Callable, TypedDict

from docx import Document
from langchain_core.messages import HumanMessage
from langgraph.graph import END, StateGraph

from app.knowledge.query import format_context, retrieve_context
from app.model_warmup import invoke_with_retry
from app.router import get_chat_model
from app.tasks.docx_writer import write_markdown

DELIVERABLES_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "data", "deliverables")
MAX_ATTEMPTS = 3

DRAFT_SYSTEM_PROMPT = (
    "You are a Mechanical Integrity approval-note drafter. You are given the transcribed "
    "text of a field inspection report and grounded SOP context retrieved from the site's "
    "procedure library.\n"
    "Rules:\n"
    "1. Draft a formal approval note summarizing the finding and the REQUIRED next action.\n"
    "2. The required action must come from the SOP context, not from the technician's own "
    "recommendation in the report — if the technician's recommendation conflicts with what "
    "the SOP actually requires, follow the SOP and explicitly note the discrepancy.\n"
    "3. Every requirement or decision you state must include an inline citation in the exact "
    "format (Page X, Section Y), matching the SOP context provided.\n"
    "4. If the SOP context does not cover the finding, state that explicitly rather than "
    "inventing a requirement."
)

REVIEW_SYSTEM_PROMPT = (
    "You are a compliance reviewer checking a drafted approval note against the same SOP "
    "context the drafter used. Check specifically:\n"
    "1. Every (Page X, Section Y) citation in the draft actually matches a requirement present "
    "in the SOP context below.\n"
    "2. The draft's required action is the one the SOP actually mandates for this finding — "
    "not merely what the technician's own notes recommended.\n"
    "3. No SOP-mandated escalation, timeframe, or reporting requirement relevant to this "
    "finding has been omitted.\n"
    "Reply with exactly one line: 'APPROVED' if the draft is fully compliant, or "
    "'REJECTED: <specific reason and the correction needed>' if not."
)


class DocumentTaskState(TypedDict):
    image_b64: str
    extracted_report: str
    context_str: str
    draft_note: str
    feedback: str
    approved: bool
    attempts: int
    docx_path: str


def ocr_node(state: DocumentTaskState) -> dict:
    model = get_chat_model("vision")
    message = HumanMessage(
        content=[
            {
                "type": "text",
                "text": (
                    "Transcribe this inspection report's content in full plain text: every "
                    "field, label, and handwritten or typed note. Preserve which value goes "
                    "with which field. Do not summarize or omit anything."
                ),
            },
            {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{state['image_b64']}"}},
        ]
    )
    extracted = invoke_with_retry(model, [message])
    return {"extracted_report": extracted}


def _ground(state: DocumentTaskState) -> str:
    candidates = retrieve_context(state["extracted_report"], top_k=6)
    return format_context(candidates)


def draft_node(state: DocumentTaskState) -> dict:
    context_str = state["context_str"] or _ground(state)
    model = get_chat_model("reasoning")
    retry_note = f"\n\nA prior draft was rejected on review: {state['feedback']}\nCorrect it." if state["feedback"] else ""
    prompt = (
        f"Inspection report (transcribed):\n{state['extracted_report']}\n\n"
        f"SOP context:\n{context_str}\n\n"
        f"Draft the approval note now.{retry_note}"
    )
    draft = invoke_with_retry(model, [HumanMessage(DRAFT_SYSTEM_PROMPT + "\n\n" + prompt)])
    return {"context_str": context_str, "draft_note": draft, "attempts": state["attempts"] + 1}


def review_node(state: DocumentTaskState) -> dict:
    model = get_chat_model("reasoning")
    prompt = (
        f"SOP context:\n{state['context_str']}\n\n"
        f"Draft approval note:\n{state['draft_note']}\n\n"
        "Review it now."
    )
    verdict = invoke_with_retry(model, [HumanMessage(REVIEW_SYSTEM_PROMPT + "\n\n" + prompt)])
    approved = verdict.strip().upper().startswith("APPROVED")
    return {"approved": approved, "feedback": "" if approved else verdict.strip()}


def route_after_review(state: DocumentTaskState) -> str:
    if state["approved"] or state["attempts"] >= MAX_ATTEMPTS:
        return "finalize"
    return "draft"


def finalize_node(state: DocumentTaskState) -> dict:
    os.makedirs(DELIVERABLES_DIR, exist_ok=True)
    filename = f"approval-note-{uuid.uuid4().hex[:8]}.docx"
    path = os.path.join(DELIVERABLES_DIR, filename)

    doc = Document()
    doc.add_heading("Mechanical Integrity Approval Note", level=1)
    if not state["approved"]:
        doc.add_paragraph(
            f"NOTE: Not fully approved after {state['attempts']} draft/review cycles — "
            f"outstanding reviewer concern: {state['feedback']}"
        ).italic = True
    doc.add_heading("Approval Note", level=2)
    write_markdown(doc, state["draft_note"])
    doc.add_heading("Source Inspection Report (Transcribed)", level=2)
    write_markdown(doc, state["extracted_report"])
    doc.save(path)

    return {"docx_path": filename}


def _build_graph():
    graph = StateGraph(DocumentTaskState)
    graph.add_node("ocr", ocr_node)
    graph.add_node("draft", draft_node)
    graph.add_node("review", review_node)
    graph.add_node("finalize", finalize_node)
    graph.set_entry_point("ocr")
    graph.add_edge("ocr", "draft")
    graph.add_edge("draft", "review")
    graph.add_conditional_edges("review", route_after_review, {"draft": "draft", "finalize": "finalize"})
    graph.add_edge("finalize", END)
    return graph.compile()


_graph = _build_graph()

_NODE_EVENTS = {
    "ocr": "ocr_done",
    "draft": "draft_done",
    "review": "review_done",
    "finalize": "done",
}


def run_document_task(image_b64: str, emit: Callable[[str, dict], None]) -> None:
    """Drives the graph via .stream(), translating each completed node's
    output into an SSE event — used by app.task_stream's background runner
    (see app/main.py's POST /tasks/document)."""
    initial: DocumentTaskState = {
        "image_b64": image_b64,
        "extracted_report": "",
        "context_str": "",
        "draft_note": "",
        "feedback": "",
        "approved": False,
        "attempts": 0,
        "docx_path": "",
    }
    final_state: dict = {}
    for step in _graph.stream(initial, stream_mode="updates"):
        for node_name, update in step.items():
            # LangGraph reports a node that returned {} as None here, not
            # {} — dict(None)/final_state.update(None) would raise.
            update = update or {}
            final_state.update(update)
            event = _NODE_EVENTS.get(node_name, node_name)
            emit(event, dict(update))
